"""ATS-safe Word renderer (python-docx).

Follows the research's parse-safety rules: single column, standard section headings,
contact info in the body (not header/footer), standard round bullets via real list
numbering (never unicode), standard font, US Letter, no tables/text-boxes/images."""
from __future__ import annotations
import io
import re
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .schema import Resume
from .render_pdf import mark_bold, skill_terms

FONT = "Calibri"   # widely available + clean; Arial/Georgia also fine


def _no_space(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0


def _heading(doc, text):
    p = doc.add_paragraph()
    _no_space(p); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11); r.font.name = FONT
    # bottom rule for scannability (paragraph border, not a table)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "999999")):
        bottom.set(qn(k), v)
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def _runs_with_bold(p, text):
    """Render **bold** spans; everything else plain. (No markdown leakage.)"""
    for i, seg in enumerate(re.split(r"(\*\*.+?\*\*)", text)):
        if not seg:
            continue
        r = p.add_run(seg[2:-2] if seg.startswith("**") and seg.endswith("**") else seg)
        r.font.name = FONT; r.font.size = Pt(10)
        if seg.startswith("**"):
            r.bold = True


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _no_space(p); p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    _runs_with_bold(p, text)


def render_docx(r: Resume, bold_terms=None) -> bytes:
    terms = list(dict.fromkeys(list(bold_terms or []) + skill_terms(r)))
    if terms:   # bold metrics + skills/JD keywords in bullets, summary, and title/stack lines
        r = copy.deepcopy(r)
        for e in r.experience:
            e.bullets = [mark_bold(b, terms) for b in e.bullets]
            e.title = mark_bold(e.title + (f"  |  {e.stack}" if e.stack else ""), terms)
            e.stack = ""   # folded into title above
        for p in r.projects:
            p.bullets = [mark_bold(b, terms) for b in p.bullets]
        if r.summary:
            r.summary = mark_bold(r.summary, terms)
    doc = Document()
    # US Letter + tight margins
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for m in ("top_margin", "bottom_margin"):
        setattr(sec, m, Inches(0.5))
    for m in ("left_margin", "right_margin"):
        setattr(sec, m, Inches(0.6))
    style = doc.styles["Normal"]; style.font.name = FONT; style.font.size = Pt(10)

    # name
    p = doc.add_paragraph(); _no_space(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(r.name); run.bold = True; run.font.size = Pt(18); run.font.name = FONT
    # contact (in body)
    if r.contact:
        p = doc.add_paragraph(); _no_space(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(r.contact); run.font.size = Pt(9.5); run.font.name = FONT
    if r.work_authorization:
        p = doc.add_paragraph(); _no_space(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(r.work_authorization); run.font.size = Pt(9); run.font.name = FONT

    if r.summary:
        _heading(doc, "Summary")
        p = doc.add_paragraph(); _no_space(p); p.paragraph_format.space_after = Pt(2)
        _runs_with_bold(p, r.summary)

    if r.experience:
        _heading(doc, "Experience")
        for e in r.experience:
            p = doc.add_paragraph(); _no_space(p)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), WD_TAB_ALIGNMENT.RIGHT)
            left = p.add_run(e.company); left.bold = True; left.font.size = Pt(10.5); left.font.name = FONT
            if e.dates:
                p.add_run("\t"); d = p.add_run(e.dates); d.font.size = Pt(9.5); d.font.name = FONT
            # title line
            t = doc.add_paragraph(); _no_space(t)
            t.paragraph_format.space_after = Pt(1)
            _runs_with_bold(t, e.title + (f"  |  {e.stack}" if e.stack else ""))
            for b in e.bullets:
                _bullet(doc, b)

    if r.projects:
        _heading(doc, "Projects")
        for pr in r.projects:
            p = doc.add_paragraph(); _no_space(p)
            run = p.add_run(pr.name + (f"  |  {pr.stack}" if pr.stack else ""))
            run.bold = True; run.font.size = Pt(10.5); run.font.name = FONT
            for b in pr.bullets:
                _bullet(doc, b)

    if r.skills:
        _heading(doc, "Skills")
        for line in r.skills:
            p = doc.add_paragraph(); _no_space(p); p.paragraph_format.space_after = Pt(1)
            if ":" in line:
                cat, items = line.split(":", 1)
                rc = p.add_run(cat.strip() + ": "); rc.bold = True; rc.font.size = Pt(10); rc.font.name = FONT
                ri = p.add_run(items.strip()); ri.font.size = Pt(10); ri.font.name = FONT
            else:
                _runs_with_bold(p, line)

    if r.education:
        _heading(doc, "Education")
        for ed in r.education:
            p = doc.add_paragraph(); _no_space(p)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), WD_TAB_ALIGNMENT.RIGHT)
            run = p.add_run(ed.school); run.bold = True; run.font.size = Pt(10.5); run.font.name = FONT
            if ed.dates:
                p.add_run("\t"); d = p.add_run(ed.dates); d.font.size = Pt(9.5); d.font.name = FONT
            if ed.degree:
                t = doc.add_paragraph(); _no_space(t)
                tr = t.add_run(ed.degree); tr.font.size = Pt(10); tr.font.name = FONT
            if ed.coursework:
                t = doc.add_paragraph(); _no_space(t)
                tr = t.add_run("Coursework: " + ed.coursework); tr.font.size = Pt(9.5); tr.font.name = FONT

    if r.achievements:
        _heading(doc, "Achievements")
        for a in r.achievements:
            _bullet(doc, a)

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()
